import json
from objets.QuestionClasse import Question

import PyPDF2
from docx import Document

def CreerObjetQuestion(path=r"code\data\questiontemp.json") -> list[Question]:
    """Permet de charger un JSON de question et de créer les objets Questions"""
    listeQuestions = []
    with open(path, "r", encoding="UTF-8") as file:
        
        data = json.load(file)
        
    for dict in data:
        listeQuestions.append(Question(dict["question"], dict["documents"]))
    
    return listeQuestions


def UpdateObjetQuestion(questions: list, documents: list) -> list:
    """Permet de mettre à jour les objets Questions"""
    documents_dict = {str(document): document for document in documents}
    
    for question in questions:
        listnouv = [
            documents_dict[docQ]
            for docQ in question.GetDocuments()
            if docQ in documents_dict
        ]
        question.SetDocument(listnouv)
    return questions



def PdfOrDocx(path : str) -> str:
    """Vérifie si le document donnée est un pdf ou un docx, il effectue la bonne fonction"""
    if path[-4:] == '.pdf':
        return PdfToString(path)
    return DocxToString(path)
    
def DocxToString(path : str) -> str:
    """Permet d'extreaire le contenue d'un DOCX"""
    doc = Document(path)
    texte = "\n".join([para.text for para in doc.paragraphs])
    return texte

def PdfToString(path : str) -> str:
    """Permet d'extreaire le contenue d'un PDF"""

    texte = ""
    
    with open(path, 'rb') as file:
        lecteur = PyPDF2.PdfReader(file)
        
        for page in lecteur.pages:
            texte += page.extract_text()
            
    return texte
    